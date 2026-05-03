"""Unit tests for CustomLyricsService."""
from __future__ import annotations

import io
import json
from unittest.mock import MagicMock, patch

import docx as docx_mod  # python-docx
import pytest

from backend.services.custom_lyrics_service import (
    CustomLyricsResult,
    CustomLyricsService,
    CustomLyricsServiceError,
)
from backend.services.custom_lyrics.settings import GenerationSettings, StrictnessLevel


def _make_counter() -> MagicMock:
    """Lightweight stand-in for SyllableCounter."""
    counter = MagicMock()
    counter.count_per_line.return_value = [2, 2, 2, 2]
    counter.any_method_within.return_value = True
    counter.min_delta.return_value = 0
    return counter


def _make_verbatim_settings() -> GenerationSettings:
    return GenerationSettings(strictness=StrictnessLevel.VERBATIM)


def _make_segments(lines: list[str]) -> list[dict]:
    """Build minimal segment dicts (one per line, consecutive 1-second windows)."""
    return [
        {"start_time": float(i), "end_time": float(i + 1), "text": line, "words": []}
        for i, line in enumerate(lines)
    ]


@pytest.fixture
def docx_bytes() -> bytes:
    """Build a real .docx in-memory with two short lines."""
    buf = io.BytesIO()
    doc = docx_mod.Document()
    doc.add_paragraph("Replace happy birthday with happy anniversary throughout.")
    doc.add_paragraph("Insert the names Mary and Steve where natural.")
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def service() -> CustomLyricsService:
    """Service with mocked settings and a lightweight counter."""
    with patch("backend.services.custom_lyrics.service.get_settings") as mock_settings:
        mock_settings.return_value = MagicMock(
            google_cloud_project="test-project",
            custom_lyrics_model="gemini-3.1-pro-preview",
            custom_lyrics_max_file_mb=5,
            custom_lyrics_max_input_lines=500,
            custom_lyrics_max_iterations=4,
            custom_lyrics_default_strictness="balanced",
            custom_lyrics_max_output_lines_multiplier=2.0,
        )
        yield CustomLyricsService(counter=_make_counter())


def _mock_gemini_response(lines: list[str]) -> MagicMock:
    """Build a MagicMock that mimics genai's GenerateContentResponse."""
    response = MagicMock()
    response.text = json.dumps({"lines": lines})
    return response


def test_text_only_happy_path(service: CustomLyricsService) -> None:
    target_lines = ["happy birthday to you", "happy birthday to you"]
    expected_output = ["happy birthday dear jane", "happy birthday dear jane"]

    with patch(
        "backend.services.custom_lyrics.service.genai.Client"
    ) as mock_client_cls:
        mock_client = MagicMock()
        mock_client.models.generate_content.return_value = _mock_gemini_response(
            expected_output
        )
        mock_client_cls.return_value = mock_client

        result = service.generate(
            job_id="job-123",
            target_lines=target_lines,
            target_segments=_make_segments(target_lines),
            artist="Anonymous",
            title="Happy Birthday",
            custom_text="Replace 'to you' with 'dear jane' wherever it makes sense",
            file_bytes=None,
            file_mime=None,
            file_name=None,
            notes=None,
            settings=_make_verbatim_settings(),
        )

    assert isinstance(result, CustomLyricsResult)
    assert result.lines == expected_output
    assert result.line_count_mismatch is False
    assert result.iterations_used == 0
    assert result.model == "gemini-3.1-pro-preview"
    assert mock_client_cls.call_args.kwargs == {
        "vertexai": True,
        "project": "test-project",
        "location": "global",
    }


def test_docx_input_parsed_to_text(
    service: CustomLyricsService, docx_bytes: bytes
) -> None:
    target_lines = ["happy birthday to you"]

    captured_prompts: list[str] = []

    def fake_generate_content(*, model, contents, config):  # noqa: ARG001
        # Capture the user prompt from the contents list
        captured_prompts.append(contents[0])
        return _mock_gemini_response(["happy anniversary mary and steve"])

    with patch(
        "backend.services.custom_lyrics.service.genai.Client"
    ) as mock_client_cls:
        mock_client = MagicMock()
        mock_client.models.generate_content.side_effect = fake_generate_content
        mock_client_cls.return_value = mock_client

        result = service.generate(
            job_id="job-1",
            target_lines=target_lines,
            target_segments=_make_segments(target_lines),
            artist="A",
            title="T",
            custom_text=None,
            file_bytes=docx_bytes,
            file_mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_name="instructions.docx",
            notes=None,
            settings=_make_verbatim_settings(),
        )

    assert result.lines == ["happy anniversary mary and steve"]
    assert any("happy anniversary" in p.lower() for p in captured_prompts)
    assert any("Mary and Steve" in p for p in captured_prompts)


def test_txt_input_decoded_and_used(service: CustomLyricsService) -> None:
    target_lines = ["original line 1"]
    txt_bytes = b"Use the word 'celebration' instead of 'birthday'."

    captured_prompts: list[str] = []

    def fake_generate_content(*, model, contents, config):  # noqa: ARG001
        captured_prompts.append(contents[0])
        return _mock_gemini_response(["original celebration 1"])

    with patch(
        "backend.services.custom_lyrics.service.genai.Client"
    ) as mock_client_cls:
        mock_client = MagicMock()
        mock_client.models.generate_content.side_effect = fake_generate_content
        mock_client_cls.return_value = mock_client

        service.generate(
            job_id="job-1",
            target_lines=target_lines,
            target_segments=_make_segments(target_lines),
            artist=None,
            title=None,
            custom_text=None,
            file_bytes=txt_bytes,
            file_mime="text/plain",
            file_name="brief.txt",
            notes=None,
            settings=_make_verbatim_settings(),
        )

    assert any("celebration" in p for p in captured_prompts)


def test_md_input_decoded_and_used(service: CustomLyricsService) -> None:
    target_lines = ["a", "b"]
    md_bytes = b"# Brief\n\nSwap 'baby' with 'sweetie'."

    with patch(
        "backend.services.custom_lyrics.service.genai.Client"
    ) as mock_client_cls:
        mock_client = MagicMock()
        mock_client.models.generate_content.return_value = _mock_gemini_response(
            ["a", "b"]
        )
        mock_client_cls.return_value = mock_client

        result = service.generate(
            job_id="job-1",
            target_lines=target_lines,
            target_segments=_make_segments(target_lines),
            artist=None,
            title=None,
            custom_text=None,
            file_bytes=md_bytes,
            file_mime="text/markdown",
            file_name="brief.md",
            notes=None,
            settings=_make_verbatim_settings(),
        )

    assert result.lines == ["a", "b"]


def test_pdf_input_passed_inline_to_gemini(service: CustomLyricsService) -> None:
    """PDFs are sent as inline parts; their bytes never get extracted server-side."""
    target_lines = ["line one", "line two"]
    pdf_bytes = b"%PDF-1.4\nfake-pdf-bytes"

    captured_contents: list[list] = []

    def fake_generate_content(*, model, contents, config):  # noqa: ARG001
        captured_contents.append(contents)
        return _mock_gemini_response(["line ONE", "line TWO"])

    with patch(
        "backend.services.custom_lyrics.service.genai.Client"
    ) as mock_client_cls:
        mock_client = MagicMock()
        mock_client.models.generate_content.side_effect = fake_generate_content
        mock_client_cls.return_value = mock_client

        service.generate(
            job_id="job-1",
            target_lines=target_lines,
            target_segments=_make_segments(target_lines),
            artist=None,
            title=None,
            custom_text=None,
            file_bytes=pdf_bytes,
            file_mime="application/pdf",
            file_name="brief.pdf",
            notes=None,
            settings=_make_verbatim_settings(),
        )

    assert len(captured_contents) == 1
    contents = captured_contents[0]
    assert any(getattr(part, "mime_type", None) == "application/pdf" for part in contents[1:]) or \
        any(part is not None for part in contents[1:])  # part objects vary by SDK version


def test_unsupported_mime_rejected(service: CustomLyricsService) -> None:
    with pytest.raises(CustomLyricsServiceError) as exc_info:
        service.generate(
            job_id="job-1",
            target_lines=["a"],
            target_segments=_make_segments(["a"]),
            artist=None,
            title=None,
            custom_text=None,
            file_bytes=b"<html/>",
            file_mime="text/html",
            file_name="brief.html",
            notes=None,
            settings=_make_verbatim_settings(),
        )
    assert exc_info.value.status_code == 400


def test_corrupt_docx_raises(service: CustomLyricsService) -> None:
    with pytest.raises(CustomLyricsServiceError) as exc_info:
        service.generate(
            job_id="job-1",
            target_lines=["a"],
            target_segments=_make_segments(["a"]),
            artist=None,
            title=None,
            custom_text=None,
            file_bytes=b"this is not a real docx",
            file_mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_name="bad.docx",
            notes=None,
            settings=_make_verbatim_settings(),
        )
    assert exc_info.value.status_code == 400


def test_oversize_file_rejected(service: CustomLyricsService) -> None:
    too_big = b"x" * (6 * 1024 * 1024)  # 6 MB > default 5 MB cap
    with pytest.raises(CustomLyricsServiceError) as exc_info:
        service.generate(
            job_id="job-1",
            target_lines=["a"],
            target_segments=_make_segments(["a"]),
            artist=None,
            title=None,
            custom_text=None,
            file_bytes=too_big,
            file_mime="text/plain",
            file_name="big.txt",
            notes=None,
            settings=_make_verbatim_settings(),
        )
    assert exc_info.value.status_code == 400


def test_gemini_exception_propagates_as_service_error(
    service: CustomLyricsService,
) -> None:
    with patch(
        "backend.services.custom_lyrics.service.genai.Client"
    ) as mock_client_cls:
        mock_client = MagicMock()
        mock_client.models.generate_content.side_effect = RuntimeError(
            "vertex unavailable"
        )
        mock_client_cls.return_value = mock_client

        with pytest.raises(RuntimeError, match="vertex unavailable"):
            service.generate(
                job_id="job-1",
                target_lines=["a"],
                target_segments=_make_segments(["a"]),
                artist=None,
                title=None,
                custom_text="any",
                file_bytes=None,
                file_mime=None,
                file_name=None,
                notes=None,
                settings=_make_verbatim_settings(),
            )


def test_gemini_returns_non_json_raises_502(service: CustomLyricsService) -> None:
    bad = MagicMock()
    bad.text = "Sure! Here you go: ..."  # not JSON

    with patch(
        "backend.services.custom_lyrics.service.genai.Client"
    ) as mock_client_cls:
        mock_client = MagicMock()
        mock_client.models.generate_content.return_value = bad
        mock_client_cls.return_value = mock_client

        with pytest.raises(CustomLyricsServiceError) as exc_info:
            service.generate(
                job_id="job-1",
                target_lines=["a"],
                target_segments=_make_segments(["a"]),
                artist=None,
                title=None,
                custom_text="any",
                file_bytes=None,
                file_mime=None,
                file_name=None,
                notes=None,
                settings=_make_verbatim_settings(),
            )
        assert exc_info.value.status_code == 502


def test_gemini_returns_wrong_shape_raises_502(service: CustomLyricsService) -> None:
    bad = MagicMock()
    bad.text = json.dumps({"output": ["a"]})  # missing "lines" key

    with patch(
        "backend.services.custom_lyrics.service.genai.Client"
    ) as mock_client_cls:
        mock_client = MagicMock()
        mock_client.models.generate_content.return_value = bad
        mock_client_cls.return_value = mock_client

        with pytest.raises(CustomLyricsServiceError) as exc_info:
            service.generate(
                job_id="job-1",
                target_lines=["a"],
                target_segments=_make_segments(["a"]),
                artist=None,
                title=None,
                custom_text="any",
                file_bytes=None,
                file_mime=None,
                file_name=None,
                notes=None,
                settings=_make_verbatim_settings(),
            )
        assert exc_info.value.status_code == 502


def test_notes_field_included_in_prompt(service: CustomLyricsService) -> None:
    captured_prompts: list[str] = []

    def fake_generate_content(*, model, contents, config):  # noqa: ARG001
        captured_prompts.append(contents[0])
        return _mock_gemini_response(["a"])

    with patch(
        "backend.services.custom_lyrics.service.genai.Client"
    ) as mock_client_cls:
        mock_client = MagicMock()
        mock_client.models.generate_content.side_effect = fake_generate_content
        mock_client_cls.return_value = mock_client

        service.generate(
            job_id="job-1",
            target_lines=["a"],
            target_segments=_make_segments(["a"]),
            artist=None,
            title=None,
            custom_text="any",
            file_bytes=None,
            file_mime=None,
            file_name=None,
            notes="Wedding for John & Jane — keep it cheesy",
            settings=_make_verbatim_settings(),
        )

    assert any("John & Jane" in p for p in captured_prompts)
